from flask import Flask, render_template, request, send_file
from docx import Document
import io
import os
import re
from datetime import datetime

app = Flask(__name__)

@app.route("/", methods=['GET'])
def home():
    return render_template("form.html")

@app.route("/generate", methods=['POST'])
def generate_doc():
    try:
        # 檢查模板是否存在
        if not os.path.exists("templates/form_template.docx"):
            return "Word模板文件不存在", 404

        doc = Document("templates/form_template.docx")
        form_data = request.form

        # 準備替換字典
        replacements = {}
        
        # 1. 處理基本字段
        for key, value in form_data.items():
            if value:  # 只處理有值的字段
                clean_key = re.sub(r'[\\_{}\s]', '', key)
                replacements[f"{{{{{clean_key}}}}}"] = value
                replacements[f"{{{{{key}}}}}"] = value

        # 2. 特殊處理空載數據
        no_load_fields = ['time', 'rpm', 'hz', 'kw', 'voltage_rs', 'voltage_st', 
                         'voltage_tr', 'current_r', 'current_s', 'current_t']
        for field in no_load_fields:
            placeholder = f"{{{{no_load_{field}}}}}"
            replacements[placeholder] = form_data.get(f"no_load_{field}", "")

        # 3. 處理加載數據 (1-5組)
        for i in range(1, 6):
            for field in ['time', 'rpm', 'hz', 'kw', 'voltage_rs', 'voltage_st', 
                         'voltage_tr', 'current_r', 'current_s', 'current_t']:
                placeholder = f"{{{{load_{field}_{i}}}}}"
                replacements[placeholder] = form_data.get(f"load_{field}_{i}", "")

        # 4. 處理反斜線變體
        additional_ph = {}
        for ph, val in replacements.items():
            if '_' in ph:
                new_ph = ph.replace('_', r'\_')
                additional_ph[new_ph] = val
        replacements.update(additional_ph)

        # 執行替換（段落+表格）
        for para in doc.paragraphs:
            para_text = para.text
            for ph, val in replacements.items():
                para_text = para_text.replace(ph, val)
            para.text = para_text
            
            for run in para.runs:
                run_text = run.text
                for ph, val in replacements.items():
                    run_text = run_text.replace(ph, val)
                run.text = run_text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    for ph, val in replacements.items():
                        cell_text = cell_text.replace(ph, val)
                    cell.text = cell_text
                    
                    for para in cell.paragraphs:
                        para_text = para.text
                        for ph, val in replacements.items():
                            para_text = para_text.replace(ph, val)
                        para.text = para_text

        # 清理未替換的佔位符
        clean_pattern = re.compile(r'\{\{[^}]*\}\}')
        for para in doc.paragraphs:
            para.text = clean_pattern.sub('', para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = clean_pattern.sub('', cell.text)

        # 生成文件
        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        filename = f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return send_file(
            output_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        app.logger.error(f"文件生成失敗: {str(e)}", exc_info=True)
        return f"文件生成失敗: {str(e)}", 500

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5000, debug=True)
